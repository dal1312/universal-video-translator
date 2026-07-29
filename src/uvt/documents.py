from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from .screen_assistant import ScreenAssistantError, extract_text


class DocumentError(RuntimeError):
    pass


MAX_DOCUMENT_CHARS = 120_000
MAX_SCANNED_PAGES = 30
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".rst",
    ".csv",
    ".json",
    ".xml",
    ".html",
    ".htm",
    ".py",
    ".js",
    ".ts",
    ".css",
    ".log",
}


@dataclass(frozen=True, slots=True)
class DocumentContext:
    title: str
    text: str
    kind: str
    pages: int = 1


def _limit(text: str) -> str:
    cleaned = text.strip()
    if len(cleaned) <= MAX_DOCUMENT_CHARS:
        return cleaned
    return (
        cleaned[:MAX_DOCUMENT_CHARS]
        + "\n\n[Documento abbreviato per il limite di contesto]"
    )


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            raise DocumentError(f"Impossibile leggere il file: {exc}") from exc
    raise DocumentError("Codifica del file di testo non riconosciuta.")


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentError(
            "Supporto DOCX non installato. Esegui INSTALL_WINDOWS.bat."
        ) from exc
    try:
        document = Document(path)
    except Exception as exc:
        raise DocumentError(f"DOCX non leggibile: {exc}") from exc
    parts = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(part for part in parts if part)


def _ocr_pdf(path: Path) -> tuple[str, int]:
    try:
        import pymupdf
        from PIL import Image
    except ImportError as exc:
        raise DocumentError(
            "Supporto PDF scansionati non installato. "
            "Esegui INSTALL_WINDOWS.bat."
        ) from exc
    try:
        document = pymupdf.open(path)
        pages = min(len(document), MAX_SCANNED_PAGES)
        texts = []
        for index in range(pages):
            page = document[index]
            pixmap = page.get_pixmap(
                matrix=pymupdf.Matrix(1.8, 1.8), alpha=False
            )
            image = Image.frombytes(
                "RGB", (pixmap.width, pixmap.height), pixmap.samples
            )
            try:
                page_text = extract_text(image)
            except ScreenAssistantError:
                page_text = ""
            if page_text:
                texts.append(f"[Pagina {index + 1}]\n{page_text}")
        return "\n\n".join(texts), len(document)
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError(
            f"OCR del documento PDF non riuscito: {exc}"
        ) from exc


def _read_pdf(path: Path) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentError(
            "Supporto PDF non installato. Esegui INSTALL_WINDOWS.bat."
        ) from exc
    try:
        reader = PdfReader(path)
        texts = []
        for index, page in enumerate(reader.pages):
            value = (page.extract_text() or "").strip()
            if value:
                texts.append(f"[Pagina {index + 1}]\n{value}")
        combined = "\n\n".join(texts)
        if len(combined.strip()) >= 40:
            return combined, len(reader.pages)
    except Exception as exc:
        raise DocumentError(f"PDF non leggibile: {exc}") from exc
    return _ocr_pdf(path)


def load_document(path: str | Path) -> DocumentContext:
    source = Path(path)
    if not source.is_file():
        raise DocumentError(f"File non trovato: {source}")
    suffix = source.suffix.casefold()
    pages = 1
    if suffix == ".pdf":
        text, pages = _read_pdf(source)
        kind = "PDF"
    elif suffix == ".docx":
        text = _read_docx(source)
        kind = "DOCX"
    elif suffix in IMAGE_SUFFIXES:
        try:
            from PIL import Image

            with Image.open(source) as image:
                text = extract_text(image.convert("RGB"))
        except ScreenAssistantError as exc:
            raise DocumentError(str(exc)) from exc
        except Exception as exc:
            raise DocumentError(f"Immagine non leggibile: {exc}") from exc
        kind = "Immagine OCR"
    elif suffix in TEXT_SUFFIXES:
        text = _read_text(source)
        kind = "Testo"
    else:
        raise DocumentError(
            f"Formato non supportato: {suffix or 'senza estensione'}"
        )
    text = _limit(text)
    if not text:
        raise DocumentError("Il documento non contiene testo rilevabile.")
    return DocumentContext(
        title=source.name,
        text=text,
        kind=kind,
        pages=pages,
    )

